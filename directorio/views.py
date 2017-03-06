#python-docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches

#template-docx
from templated_docs import fill_template
from templated_docs.http import FileResponse

#modelos
from directorio.models import Directorio

#django
from django.http import HttpResponse, HttpResponseRedirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Q

#Modelos
from directorio.models import Directorio, Obsequio, Acuse, Historial

#Formularios
from directorio.forms import DirectorioForm, ObsequioForm

# CRUD Directorio
@login_required
def home(request):
    """ 
        se requeire sesion iniciada

        Se realiza la consulta de todo el directorio(todas las personas que conforman el directorio)
        , si se llega a este punto con un parametro de busuqeda, se filtra la busqeda total, para mostrar
        las coincidencias(contenidas en la columnas profesion, nombre y cargo).
    """
    instancias = Directorio.objects.all().order_by("id")

    query = request.GET.get("q")
    if query:
        instancias = instancias.filter(Q(profesion__icontains = query)|
                                        Q(nombre__icontains = query) |
                                        Q(cargo__icontains = query) )

    contexto ={
        "directorio" : instancias,
    }
    return render(request,'home.html', contexto)

@login_required
def create_directorio(request):
    """
        se requeire sesion iniciada

        Creacion del formulario, para el registro de una nueva persona,
        se valida la informacion y se envia un mensaje de alerta.
    """
    title = ""
    form = DirectorioForm(request.POST or None)

    if form.is_valid():
        instancia = form.save(commit = False)
        instancia.user = request.user
        instancia.save()
        status = instancia.status
        messages.success(request,"!Registrado Correctamente¡")
        return HttpResponseRedirect(instancia.get_detail_path())

    contexto = {
        "title" : title,
        "form"  : form
    }
    return render(request,"form_dir.html",contexto)

@login_required
def milista_dir(request):
    """
    se requeire sesion iniciada

    se realiza la consulta de todas las personas que conforman la lista que crea el usuario, en su sesión
    , si se llega a este punto con un parametro de busqeuda, se filtra la lista y se muestran las
    coincidencias.
    """
    lista = Directorio.objects.filter(user_id__exact= request.user.id).order_by("id")

    query = request.GET.get("q")

    if query:
        lista = lista.filter(Q(profesion__icontains = query)|
                            Q(nombre__icontains = query) |
                            Q(cargo__icontains = query))
    contexto = {
        'lista' : lista
    }
    return render(request,'milista.html',contexto)

@login_required
def edit_directorio(request, id = None):
    """
        se requeire sesion iniciada

        esta vsta pide un ID de parametro, para realizar la busqueda, de la persona
        a ediar, posteriormente se crea el formulario con los datos a mostrar, se valida la informacion
        y se envia un mensjae de alerta.
    """
    persona = get_object_or_404(Directorio,id = id)
    pstatus = persona.status
    title   = "Editando a: \r %s"%persona.nombre
    form    = DirectorioForm(request.POST or None,instance = persona)
    
    if form.is_valid():
        instancia = form.save(commit = False)
        if instancia.user_id != request.user.id:
            instancia.modificado = request.user.username

        instancia.save()
        messages.success(request,"!Modificado correctamente¡")
        return HttpResponseRedirect(instancia.get_detail_path())
    contexto = {
        "title" : title,
        "form" : form
    }

    return render(request,"form_dir.html",contexto)

@login_required
def detail_directorio(request, id = None):
    """
        se reqeuire sesion iniciada

        Requeire un parametro ID para realizar la busqeuda de los detalles a mostrar de la
        persona
    """
    instancia = get_object_or_404(Directorio, id = id)
    contexto = {
        "title": "Detalles",
        "persona" :instancia,
    }

    return render(request,"detalle_dir.html",contexto)

@login_required
def confirm_delete_directorio(request,id = None):
    """
        se requiere sesion iniciada

        REquiere un parametro ID para realizar de la busqeda de la persona
        y mostrar los detalles antes de eliminar
    """
    instancia = get_object_or_404(Directorio, id = id)
    contexto = {
        "persona" : instancia,
    }
    return render(request,'confirm_delete.html',contexto)

@login_required
def delete_directorio(request, id=None ):
    """
        Se requiere sesion inicida

        requeire un parametro ID para realizar la eliminacion de
        de la persona seleccionada
    """
    instancia = get_object_or_404(Directorio, id= id)
    instancia.delete()
    return redirect('directorio:eliminado')

@login_required
def load_detail(request, id = None):
    """
    requeire sesion iniciada

    requiere un parametro ID, para mostrar los datos de una persona,
    contenido que se mostrara por AJAX jquery
    """
    instancia = get_object_or_404(Directorio, id = id)
    contexto = {
        "persona" : instancia
    }
    return render(request,"load_detail.html",contexto)

@login_required
def informacion(request):
    """ 
    No esta en funcionamiento en la aplicación
    """
    instancia  = Obsequio.objects.get(default = True)
    entregados = instancia.entregado #Directorio.objects.filter(status__exact = 3).count()
    existencia = instancia.existencia
    contexto   = {
        "obsequio" : instancia,
        "entregados" : entregados,
        "existencia"    : existencia
    }
    return render(request,"informacion.html",contexto)

@login_required
def informacion2(request):
    """
        requiere sesion iniciada

        Muestra el total de personas autorizadas y entregadas
    """
    instancia  = Obsequio.objects.get(default = True)
    entregados = Directorio.objects.filter(status__exact = 3).count()
    autorizados = Directorio.objects.filter(status__exact = 2).count()

    contexto= {
        "obsequio" : instancia,
        "entregados" : entregados,
        "autorizados" : autorizados
    }

    return render(request,"informacion.html",contexto)

@login_required
def eliminado(request):
    """
    Requiere secion iniciada

    Simplemente renderiza un template con mesaje de eliminado
    """
    return render(request,"eliminado.html")


def agregar_mi_lista(request,id = None):
    """
    Requiere un parametro ID, obtiene la persona indicada, para poder editar el user_id
    el cual sirve para que el usuario, pueda crear un alista con la cual trabajara. Esta accion se realiza a travez
    de AJAX jquery.
    """
    instancia = get_object_or_404(Directorio, id = id)
    instancia.user_id = request.user.id
    instancia.lista = True
    instancia.save()

    return HttpResponse("se añadio correctamente")

def quitar_mi_lista(request, id = None):
    """
    reqeuire un parametro ID, para relizar la busqueda de la persona,
    y poder modificar el user_id. El usuario podra reducir su lista de trabajo. Esta accion se
    realiza atravez de AJAX Jquery.
    """
    instancia = get_object_or_404(Directorio, id = id)
    instancia.lista = False
    instancia.user_id = None
    instancia.save()

    return HttpResponse("!Ya no estara en tu lista¡")

#Configuracion y perfil de usuario

@login_required
def config_user(request):
    return render(request,"configuracion_user.html")

# CRUD de Acuse
@login_required
def acuses_lista(request):
    """
        requiere sesion iniciada
        Realiza la busqueda de todos los acuses, y los muestra en lista
    """
    instancias = Acuse.objects.all()
    contexto = {
        "acuses" : instancias
    }
    return render(request,"acuses.html",contexto)

@login_required
def ver_acuse(request):
    """
    No esta en funcionamiento
    """
    instancia = Acuse.objects.get(default=True)
    contexto = {
        "acuse": instancia,
        "llamado":"ver" 
    }
    return render(request,"editar_acuse.html",contexto)

@login_required
def nuevo_acuse(request):
    """
    requiere secion inicida

    Crea un nuevo acuse, obteniendo los datos de formulario
    """
    acuse = Acuse()
    if request.POST:
        acuse.alias = request.POST.get("alias")
        acuse.contenido = request.POST.get("contenido")
        acuse.save()
        return redirect("directorio:acuses_lista")
    return render(request,"new_acuse.html")
@login_required
def editar_acuse(request, id = None):
    """
    requiere secion inicida

    REquiere un parametro ID para realizar la busqeuda de acuse,
    y poder editar los datos, obtenidos del fromulario
    """
    instancia = Acuse.objects.get(id = id)
    if request.POST:
        if request.POST.get("contenido") != "":
            instancia.contenido = request.POST.get("contenido")
            instancia.alias = request.POST.get("alias")
            instancia.save()
    
    contexto = {
        "acuse":instancia,
        "llamado" : "editar"
    }
    return render(request,"editar_acuse.html",contexto)

# vista de reportes
@login_required
def inicio_reportes(request):
    """
        crea una lista segun el boton indicado, relacionde de todos los entregados, relacion de los autorizados
        ,relacion de los no autorizados, relacion de cada acuse.
    """

    if request.POST:
        if request.POST.get("tipo") == "entregado":
            """
                Crea una relacion de las personas a las que se le entrego obsequio
            """
            instancias = Directorio.objects.filter(status__exact = 3)
            documento = Document()
            documento.add_paragraph("Relacion de Entregados")
            return crear_lista(documento,instancias)

        elif request.POST.get("tipo") == "pendiente":
            """
                crea la relacion de la sspersonas que solo estan autorizadas pero no sel eha entregado aobsequio
            """
            documento = Document()
            documento.add_paragraph("Relacacion de Autorizados")
            instancias = Directorio.objects.filter(status__exact = 2)
            return crear_lista(documento,instancias)
        elif request.POST.get("tipo") == "noautorizado":
            """
                crea la relacion de las personas que no se les autorio obsequio
            """
            documento =Document()
            documento.add_paragraph("Relacion de los No Autoriados")
            instancias = Directorio.objects.filter(status__exact = 1)
            return crear_lista(documento,instancias)            

    return render(request,"inicio_reportes.html")
# funcion para crear lista
def crear_lista(documento,instancias):

    for persona in instancias:
        documento.add_paragraph("%s%s"%(persona.profesion,persona.nombre), style='ListNumber')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=relacion_entregados.docx'
    documento.save(response)

    return response

# Vistas que generan un archivo de WORD
@login_required
def seleccionar_acuses_mi_lista(request):
    """
    require sescion iniciada
    Lista creada por el usuario.
    muestra la lista de personas que han sido autorizadas y entregadas
    en el template hay tres opciones
    acuses:
        esta opcion crea un archivo de word con los acuses de las personas seleccioandas
    etiquetas:
        esta opcion crea unarchivo de word con las etiquetas de las personas seleccioandas
    lista:
        esta opción genera un archivo de word con el nombre de las personas seleccionadas
    """
    instancias = Directorio.objects.filter(user_id__exact = request.user.id).exclude(status__exact=1)

    mensaje = "Selecciona los acuses de tu lista"

    invocador = "milista"

    if request.POST:
        if request.POST.get("tipo") == "acus":
            acuses = request.POST.getlist("acuses")
            document = Document()

            for x in acuses:
                persona = get_object_or_404(Directorio, id= x)
                acuse(document,persona)
                document.add_page_break()

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=acuses.docx'
            document.save(response)
            return response

        elif request.POST.get("tipo") == "eti":
            etiquetas = request.POST.getlist("acuses")
            personas = []
            num_etiquetas = len(etiquetas)
            print(num_etiquetas)
            print(etiquetas)

            for i in etiquetas:
                personas.append(Directorio.objects.get(id = i))

            blanco = Directorio()
            blanco.profesion = " "
            blanco.nombre = " "
            blanco.cargo = " "
            blanco.direccion = " "
            blanco.pareja = None
            con1 = []
            con2 = []
            rango  = 0

            if num_etiquetas != 1:
                x = 0
                while x < num_etiquetas:
                    if (x%2==0):
                        con1.append(personas[x])
                    else:
                        con2.append(personas[x])
                    x =x+1
                if len(con1) != len(con2):
                    if len(con1) > len(con2):
                        con2.append(blanco)
                    elif len(con2) > len(con1):
                        con1.append(blanco)
            else:
                con1.append(personas[0])
                con2.append(blanco)

            context = {
                'rango': range(len(con1)),
                # 'personas': personas,
                'con1' : con1,
                'con2'  : con2
            }

            filename = fill_template('etiquetas.odt', context, output_format='docx')
            visible_filename = 'mis_etiquetas.docx'

            return FileResponse(filename, visible_filename)

        elif request.POST.get("tipo") == "lista":
            documento = Document()  
            
            

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=acuse.docx'
            documento.save(response)

            return response


    contexto = {
        "instancias" : instancias,
        "mensaje" : mensaje,
        "invocador" : invocador
    }
    return render(request,"multiple_acuses.html",contexto)

@login_required
def seleccionar_acuses_directorio(request):
    """
    require sescion iniciada

    Lista general de personas
    muestra la lista de personas que han sido autorizadas y entregadas
    en el template hay tres opciones
    acuses:
        esta opcion crea un archivo de word con los acuses de las personas seleccioandas
    etiquetas:
        esta opcion crea unarchivo de word con las etiquetas de las personas seleccioandas
    lista:
        esta opción genera un archivo de word con el nombre de las personas seleccionadas
    """
    instancias = Directorio.objects.all().exclude(status__exact = 1)

    mensaje = "*Aqui podras seleccionar los acuses que han sido autorizados en todo el directorio"

    invocador = "directorio"

    if request.POST:
        if request.POST.get("tipo") == "acus":
            acuses = request.POST.getlist("acuses")

            document = Document()

            for x in acuses:
                # print(x)
                persona = get_object_or_404(Directorio, id= x)
                acuse(document,persona)
                document.add_page_break()

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=acuses.docx'
            document.save(response)

            return response

        elif request.POST.get("tipo") == "eti":
            etiquetas = request.POST.getlist("acuses")
            personas = []
            num_etiquetas = len(etiquetas)
            # print(num_etiquetas)
            # print(etiquetas)

            for i in etiquetas:
                personas.append(Directorio.objects.get(id = i))

            blanco = Directorio()
            blanco.profesion = " "
            blanco.nombre = " "
            blanco.cargo = " "
            blanco.direccion = " "
            blanco.pareja = None
            con1 = []
            con2 = []
            rango  = 0

            if num_etiquetas != 1:
                x = 0
                while x < num_etiquetas:
                    if (x%2==0):
                        con1.append(personas[x])
                    else:
                        con2.append(personas[x])
                    x =x+1
                if len(con1) != len(con2):
                    if len(con1) > len(con2):
                        con2.append(blanco)
                    elif len(con2) > len(con1):
                        con1.append(blanco)
            else:
                con1.append(personas[0])
                con2.append(blanco)

            context = {
                'rango': range(len(con1)),
                # 'personas': personas,
                'con1' : con1,
                'con2'  : con2
            }

            filename = fill_template('etiquetas.odt', context, output_format='docx')
            visible_filename = 'etiquetas.docx'

            return FileResponse(filename, visible_filename)  

        elif request.POST.get("tipo") == "lista":
            lista = request.POST.getlist("acuses")
            document = Document()
            personas = []

            for y in lista:
                personas.append(Directorio.objects.get(id = y))
            
            for x in personas:
                document.add_paragraph(x.profesion+" "+x.nombre, style='ListNumber')

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=lista.docx'
            document.save(response)

            return response


    contexto = {
        "instancias" : instancias,
        "mensaje"   : mensaje,
        "invocador" : invocador
    }

    return render(request,"multiple_acuses.html",contexto)


def un_acuse(request,id = None):
    """
    requiere un prametro ID para realizar busqeuda de la persona,
    y poder crear el acuse con los datos de la misma.
    """
    persona = get_object_or_404(Directorio, id = id)
    documento = Document()  
    
    acuse(documento,persona)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=acuse.docx'
    documento.save(response)

    return response

def crear_mi_lista(request):
    """
        Esta opcion crea la lista de las personas, de cda usuario
    """
    instancias = Directorio.objects.filter(user_id__exact = request.user.id)
    instancias = instancias.filter(status__exact = 2)
    document =Document()

    for persona in instancias:
        document.add_paragraph("%s%s"%(persona.profesion,persona.nombre))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=mi_lista.docx'
    document.save(response)

    return response


def crear_mis_acuses(request):
    """
    Esta opcion crea los acuses de las personas autorizadas en la lista de cada usuario
    """
    mis_acuses = Directorio.objects.filter(user_id__exact=request.user.id).filter(status__exact = 2)

    document = Document()

    for persona in mis_acuses:
        acuse(document,persona)
        document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=mis_acuses.docx'
    document.save(response)
    return response

def crear_lista_general(request):
    """
        esta opcion crea la lista general en un archiv de word solo nombres
    """
    instancias = Directorio.objects.all()
    document =Document()

    for persona in instancias:
        document.add_paragraph("%s%s"%(persona.profesion,persona.nombre))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=directorio.docx'
    document.save(response)

    return response

def acuses_generales(request):

    """
    esta opcion crea un archivo de word con todos los acuses que se han autorizado en el directorio
    """
    acuses = Directorio.objects.all().filter(status__exact=2)

    document = Document()

    for persona in acuses:
        acuse(document,persona)
        document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=acuses.docx'
    document.save(response)
    return response 

def mis_etiquetas(request):
    """
    
    """
    personas = Directorio.objects.filter(user_id__exact = request.user.id).filter(status__exact = 2)
    num_etiquetas = Directorio.objects.filter(user_id__exact = request.user.id).filter(status__exact = 2).count()
    blanco = Directorio()
    blanco.profesion = " "
    blanco.nombre = " "
    blanco.cargo = " "
    blanco.direccion = " "
    blanco.pareja = None
    con1 = []
    con2 = []
    rango  = 0

    if num_etiquetas != 1:
        x = 0
        while x < num_etiquetas:
            if (x%2==0):
                con1.append(personas[x])
            else:
                con2.append(personas[x])
            x =x+1
        if len(con1) != len(con2):
            if len(con1) > len(con2):
                con2.append(blanco)
            elif len(con2) > len(con1):
                con1.append(blanco)
    else:
        con1.append(personas[0])
        con2.append(blanco)

    context = {
        'rango': range(len(con1)),
        # 'personas': personas,
        'con1' : con1,
        'con2'  : con2
    }

    filename = fill_template('etiquetas.odt', context, output_format='docx')
    visible_filename = 'mis_etiquetas.docx'

    return FileResponse(filename, visible_filename)

def una_etiqueta(request,id = None):
    instancia = get_object_or_404(Directorio, id = id)
    blanco = Directorio()
    blanco.profesion = " "
    blanco.nombre = " "
    blanco.cargo = " "
    blanco.cireccion = " "
    blanco.pareja = None
    con1 = []
    con2 = []

    con1.append(instancia)
    con2.append(blanco)

    contexto = {
        "rango" : range(1),
        "con1" : con1,
        "con2" : con2
    }

    filename = fill_template('etiquetas.odt', contexto, output_format='docx')
    visible_filename = 'etiqueta.docx'

    return FileResponse(filename, visible_filename)

def all_etiquetas(request):
    # contando los autorizados
    num_personas = Directorio.objects.all().filter(status__exact="2").count()
    # obteniendo los acuses autorizados
    personas = Directorio.objects.all().filter(status__exact="2")

    # creando directorio falso para completar
    blanco = Directorio()
    blanco.profesion = " "
    blanco.nombre = " "
    blanco.cargo = " "
    blanco.direccion = " "
    blanco.pareja = None
    con1 = []
    con2 = []
    rango  = 0

    if num_personas != 1:
        x = 0
        while x < num_personas:
            if (x%2==0):
                con1.append(personas[x])
            else:
                con2.append(personas[x])
            x =x+1
        if len(con1) != len(con2):
            if len(con1) > len(con2):
                con2.append(blanco)
            elif len(con2) > len(con1):
                con1.append(blanco)
    else:
        con1.append(personas[0])
        con2.append(blanco)

    context = {
        'rango': range(len(con1)),
        # 'personas': personas,
        'con1' : con1,
        'con2'  : con2
    }

    filename = fill_template('etiquetas.odt', context, output_format='docx')
    visible_filename = 'todas_las_etiquetas.docx'

    return FileResponse(filename, visible_filename)


#modulo de obsequios
def control_obsequios(request):
    instancias = Obsequio.objects.all()
    contexto = {
        "obsequios": instancias
    }
    return render(request,"control_obsequios.html",contexto)

def new_obsequio(request):
    form = ObsequioForm(request.POST or None)

    if form.is_valid():
        instancia = form.save(commit =False)
        instancia.save()
        return redirect("directorio:control_obsequios")
    contexto = {
        "form": form,
    }

    return render(request,"new_obsequio.html",contexto)

def editar_obsequio(request, id = None):
    obsequio = get_object_or_404(Obsequio, id = id)

    form = ObsequioForm(request.POST or None, instance = obsequio)

    if form.is_valid():
        instancia = form.save(commit=False)
        instancia.save()

        return redirect("directorio:control_obsequios")
    contexto = {
        "form" : form
    }

    return render(request,"new_obsequio.html",contexto)

def agregar_entregados(request):
    cantidad = int(request.POST.get("cantidad"))
    if request.POST:
        instancia = Obsequio.objects.get(id = request.POST.get("libro"))
        instancia.entregado = instancia.entregado+cantidad
        instancia.existencia = instancia.existencia-cantidad
        instancia.save()

        instancias = Obsequio.objects.all()
        contexto = {
            "obsequios": instancias
        }
        return render(request,"tabla_obsequios.html",contexto)
    return raise404



# cuerpo del aacuse, recibiendo por parametro el objeto document y los datos de la persona
def acuse(document, persona):
    instancia = Acuse.objects.get(default = True)
    acuse_contenido = instancia.contenido
    x = 1;
    while x<=4:
        document.add_paragraph()
        x = x+1
    p1 = document.add_paragraph()
    p1Format = p1.paragraph_format
    p1Format.space_before = Pt(0)
    p1Format.space_after = Pt(0)
    nombre = "%s %s"%(persona.profesion,persona.nombre)
    f = p1.add_run(nombre).font
    f.name = 'Arial'
    f.bold = True
    f.size = Pt(12)
    if persona.pareja:
        strpareja = "\ny %s"%persona.pareja
        fpareja = p1.add_run(strpareja).font
        fpareja.name = 'Arial'
        fpareja.bold = True
        fpareja.size = Pt(12)

    if persona.cargo:
        cargo = persona.cargo
        cargo = cargo.replace('\r','')
        p2 = document.add_paragraph()
        p2Format = p2.paragraph_format
        p2Format.space_before = Pt(0)
        p2Format.space_after = Pt(0)
        f2 = p2.add_run("%s"%cargo).font
        f2.name = "Arial"
        f2.size = Pt(12)
    if persona.direccion:
        direccion = persona.direccion
        direccion = direccion.replace('\r','')
        p3 = document.add_paragraph()
        p3Format = p3.paragraph_format
        p3Format.space_before = Pt(0)
        p3Format.space_after = Pt(0)
        f3 = p3.add_run("%s"%direccion).font
        f3.name = "Arial"
        f3.size = Pt(12)

    x = 1
    while x<=2:
        document.add_paragraph()
        x = x+1
    
    p4 =document.add_paragraph()
    p4Format = p4.paragraph_format
    p4Format.space_before = Pt(0)
    p4Format.space_after = Pt(0)
    p4Format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4Format.first_line_indent = Inches(0.50)
    f4 = p4.add_run(persona.acuse.contenido).font
    f4.name = 'Arial'
    f4.size = Pt(12)

    x = 1
    while x <=3:
        document.add_paragraph()
        x = x+1
    p5 =document.add_paragraph()
    p5Format = p5.paragraph_format
    p5Format.space_before = Pt(0)
    p5Format.space_after = Pt(0)
    p5Format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f5 = p5.add_run('__________________________________').font
    f5.name = 'Arial'
    f5.size = Pt(12)    

    p6 =document.add_paragraph()
    p6Format = p6.paragraph_format
    p6Format.space_before = Pt(0)
    p6Format.space_after = Pt(0)
    p6Format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f6 = p6.add_run('Nombre, firma y fecha').font
    f6.name = 'Arial'
    f6.size = Pt(12)
